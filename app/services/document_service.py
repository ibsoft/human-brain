from datetime import datetime
from pathlib import Path

from flask import current_app
from werkzeug.utils import secure_filename

from app.extensions import db
from app.models import Memory, MemoryAsset
from app.services.asset_vector_service import AssetVectorService
from app.services.memory_service import MemoryService, asset_url
from app.utils.hash import sha256_text


TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".json", ".csv", ".log", ".yaml", ".yml", ".py", ".js", ".html", ".css", ".xml"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tif", ".tiff"}
AUDIO_EXTENSIONS = {".mp3", ".wav", ".flac", ".m4a", ".aac", ".ogg", ".oga", ".opus", ".wma", ".aiff", ".aif"}


class DocumentIngestionService:
    def ingest_uploads(self, uploads, base_payload, actor_type="user", actor_id=None, mode="all", chunk_size=4000):
        mode = self._normalize_mode(mode)
        created = []
        for upload in uploads:
            if not upload or not upload.filename:
                continue
            stored_path = self._save(upload)
            payloads = self._payloads_for_file(stored_path, upload.mimetype, base_payload, mode, chunk_size, upload.filename)
            created_new_asset = False
            for payload in payloads:
                existing = self._existing_uploaded_image(payload)
                if existing:
                    existing.trust_score = min(1.0, existing.trust_score + 0.05)
                    existing.updated_at = datetime.utcnow()
                    try:
                        from app.services.audit_service import AuditService

                        AuditService.log("memory.duplicate_seen", actor_type, actor_id, existing.workspace_id, "memory", existing.id)
                    except Exception:
                        current_app.logger.exception("Could not audit duplicate uploaded image")
                    db.session.commit()
                    created.append(existing)
                    continue
                memory, _ = MemoryService().add_memory(payload, actor_type=actor_type, actor_id=actor_id)
                asset = self._create_asset(memory, stored_path, upload.filename, upload.mimetype, payload.get("_asset_type", "document"), payload.get("_asset_text"))
                self._attach_asset_url(memory, asset)
                created.append(memory)
                created_new_asset = True
            if not created_new_asset:
                self._remove_saved_duplicate(stored_path)
        return created

    def replace_memory_asset(self, memory, upload, actor_type="agent", actor_id=None, title=None, tags=None, mode="full", chunk_size=4000):
        if not upload or not upload.filename:
            raise ValueError("Upload one replacement file")
        asset = MemoryAsset.query.filter_by(memory_id=memory.id).order_by(MemoryAsset.id.asc()).first()
        if not asset:
            raise ValueError("Memory has no uploaded asset to replace")
        mode = self._normalize_mode(mode)
        if mode == "chunks":
            raise ValueError("Replacing an asset updates one memory; upload chunked documents as new memories")
        old_path = Path(asset.stored_path)
        stored_path = self._save(upload)
        base_payload = {
            "agent_id": memory.agent_id,
            "workspace_id": memory.workspace_id,
            "title": title or memory.title,
            "memory_type": memory.memory_type,
            "tags": tags if tags is not None else list(memory.tags or []),
            "importance_score": memory.importance_score,
            "trust_score": memory.trust_score,
            "sensitivity_level": memory.sensitivity_level,
            "visibility": memory.visibility,
            "confirmed": memory.confirmed,
            "source": memory.source or "upload",
            "storage_reason": f"Replaced uploaded file with {stored_path.name}.",
        }
        payload = self._payloads_for_file(stored_path, upload.mimetype, base_payload, "full", chunk_size, upload.filename)[0]
        memory.title = payload.get("title") or memory.title
        memory.content = payload["content"]
        memory.summary = payload.get("summary")
        memory.memory_type = payload.get("memory_type") or memory.memory_type
        memory.tags = payload.get("tags") or []
        memory.storage_reason = payload.get("storage_reason") or base_payload["storage_reason"]
        memory.content_hash = sha256_text(memory.content.strip())
        asset_type = payload.get("_asset_type", "document")
        vector, vector_hash, metadata = self._asset_vector_for(asset_type, stored_path, upload.mimetype, payload.get("_asset_text"), memory.content)
        asset.asset_type = asset_type
        asset.original_filename = upload.filename
        asset.stored_path = str(stored_path)
        asset.content_type = upload.mimetype
        asset.vector_hash = vector_hash
        asset.vector_dim = len(vector) if vector else None
        asset.vector = vector
        asset.asset_metadata = metadata
        db.session.commit()
        self._attach_asset_url(memory, asset)
        try:
            if old_path.exists() and old_path != stored_path:
                old_path.unlink()
        except OSError:
            current_app.logger.warning("Could not remove replaced asset file: %s", old_path)
        try:
            from app.services.audit_service import AuditService

            AuditService.log("memory.asset_replaced", actor_type, actor_id, memory.workspace_id, "memory", memory.id)
            db.session.commit()
        except Exception:
            current_app.logger.exception("Could not audit uploaded asset replacement")
        try:
            from app.services.correlation_service import CorrelationService

            CorrelationService().correlate_memory(memory)
        except Exception:
            current_app.logger.exception("Could not correlate replaced uploaded asset memory")
        return memory, asset

    def _save(self, upload):
        upload_dir = Path(current_app.config["MEMORY_UPLOAD_DIR"])
        upload_dir.mkdir(parents=True, exist_ok=True)
        filename = secure_filename(upload.filename) or "upload.bin"
        stamped = f"{datetime.utcnow().strftime('%Y%m%d%H%M%S%f')}_{filename}"
        path = upload_dir / stamped
        upload.save(path)
        return path

    def _payloads_for_file(self, path, mimetype, base_payload, mode, chunk_size, original_filename=None):
        ext = path.suffix.lower()
        display_name = original_filename or path.name
        tags = list(base_payload.get("tags") or [])
        tags.extend(["upload", ext.lstrip(".") or "file"])
        common = {
            **base_payload,
            "tags": sorted(set(tags)),
            "source": "upload",
            "storage_reason": f"Stored from uploaded file {path.name}.",
        }
        if ext in IMAGE_EXTENSIONS or str(mimetype).startswith("image/"):
            vector, vector_hash, metadata = AssetVectorService().image_vector(path)
            content = (
                f"Uploaded image: {display_name}\n"
                "Asset URL: pending\n"
                f"MIME type: {mimetype or 'unknown'}\n"
                f"Image vector: {metadata.get('vector_kind')} {vector_hash}\n"
                f"Visual profile: {metadata.get('width')}x{metadata.get('height')}, dominant color {metadata.get('dominant_color')}.\n"
                "Image bytes are available through the tokenized asset URL. Visual fingerprint metadata is stored for correlation."
            )
            image_tags = sorted(set(common["tags"] + ["image", "visual", metadata.get("dominant_color", "")]))
            return [
                {
                    **common,
                    "title": common.get("title") or display_name,
                    "content": content,
                    "summary": f"Uploaded image with {metadata.get('vector_kind')} vector.",
                    "memory_type": "vision",
                    "tags": [tag for tag in image_tags if tag],
                    "_asset_type": "image",
                    "_asset_vector": vector,
                    "_asset_hash": vector_hash,
                    "_asset_metadata": metadata,
                }
            ]

        if ext in AUDIO_EXTENSIONS or str(mimetype).startswith("audio/"):
            vector, vector_hash, metadata = AssetVectorService().file_vector(path)
            size_bytes = path.stat().st_size if path.exists() else 0
            content = (
                f"Uploaded audio/music file: {display_name}\n"
                "Asset URL: pending\n"
                f"MIME type: {mimetype or 'unknown'}\n"
                f"Size: {size_bytes} bytes\n"
                f"Audio fingerprint: {metadata.get('vector_kind')} {vector_hash}\n"
                "Audio bytes are available through the tokenized asset URL."
            )
            audio_tags = sorted(set(common["tags"] + ["audio", "music"]))
            return [
                {
                    **common,
                    "title": common.get("title") or display_name,
                    "content": content,
                    "summary": "Uploaded audio/music file stored as a memory asset.",
                    "memory_type": common.get("memory_type") or "media",
                    "tags": audio_tags,
                    "_asset_type": "audio",
                    "_asset_vector": vector,
                    "_asset_hash": vector_hash,
                    "_asset_metadata": metadata,
                }
            ]

        text = self._extract_text(path, mimetype)
        vector, vector_hash, metadata = AssetVectorService().document_vector(text, path)
        common["_asset_type"] = "document"
        common["_asset_text"] = text
        common["_asset_vector"] = vector
        common["_asset_hash"] = vector_hash
        common["_asset_metadata"] = metadata
        if mode == "chunks":
            chunks = self._chunk_text(text, max(500, int(chunk_size or 4000)))
            return [
                {
                    **common,
                    "title": f"{common.get('title') or display_name} - chunk {index + 1}",
                    "content": chunk,
                    "summary": f"Chunk {index + 1} from uploaded file {display_name}.",
                }
                for index, chunk in enumerate(chunks)
            ]
        return [{**common, "title": common.get("title") or display_name, "content": text}]

    def _normalize_mode(self, mode):
        value = str(mode or "full").strip().lower()
        if value in {"all", "full", "whole", "single"}:
            return "full"
        if value in {"chunks", "chunk", "chunked"}:
            return "chunks"
        return "full"

    def _create_asset(self, memory, path, original_filename, mimetype, asset_type, asset_text=None):
        vector, vector_hash, metadata = self._asset_vector_for(asset_type, path, mimetype, asset_text, memory.content)
        db.session.add(
            asset := MemoryAsset(
                memory_id=memory.id,
                workspace_id=memory.workspace_id,
                asset_type=asset_type,
                original_filename=original_filename,
                stored_path=str(path),
                content_type=mimetype,
                vector_hash=vector_hash,
                vector_dim=len(vector) if vector else None,
                vector=vector,
                asset_metadata=metadata,
            )
        )
        db.session.commit()
        try:
            from app.services.correlation_service import CorrelationService

            CorrelationService().correlate_memory(memory)
        except Exception:
            current_app.logger.exception("Could not correlate uploaded asset memory")
        return asset

    def _existing_uploaded_image(self, payload):
        if payload.get("_asset_type") != "image" or not payload.get("_asset_hash"):
            return None
        asset = (
            MemoryAsset.query.join(Memory, Memory.id == MemoryAsset.memory_id)
            .filter(
                Memory.workspace_id == payload["workspace_id"],
                Memory.deleted_at.is_(None),
                Memory.archived.is_(False),
                MemoryAsset.asset_type == "image",
                MemoryAsset.vector_hash == payload["_asset_hash"],
            )
            .order_by(MemoryAsset.created_at.asc())
            .first()
        )
        return db.session.get(Memory, asset.memory_id) if asset else None

    def _remove_saved_duplicate(self, path):
        try:
            path.unlink(missing_ok=True)
        except OSError:
            current_app.logger.warning("Could not remove duplicate uploaded asset file: %s", path)

    def _asset_vector_for(self, asset_type, path, mimetype, asset_text=None, fallback_text=""):
        if asset_type == "image":
            return AssetVectorService().image_vector(path)
        if asset_type == "audio":
            return AssetVectorService().file_vector(path)
        return AssetVectorService().document_vector(asset_text or fallback_text, path)

    def _attach_asset_url(self, memory, asset):
        url = asset_url(asset, external=True)
        if "Asset URL: pending" in memory.content:
            memory.content = memory.content.replace("Asset URL: pending", f"Asset URL: {url}")
        elif "Asset URL:" not in memory.content:
            memory.content = f"{memory.content.rstrip()}\n\nAsset URL: {url}"
        memory.content_hash = sha256_text(memory.content.strip())
        db.session.commit()
        try:
            service = MemoryService()
            service.faiss.upsert_memory(memory, service.embedding_service)
        except Exception:
            current_app.logger.exception("Could not refresh uploaded asset memory vector after asset URL attach")

    def _extract_text(self, path, mimetype):
        ext = path.suffix.lower()
        if ext in TEXT_EXTENSIONS or str(mimetype).startswith("text/"):
            return path.read_text(encoding="utf-8", errors="replace")
        if ext == ".pdf":
            return self._extract_pdf(path)
        if ext == ".docx":
            return self._extract_docx(path)
        if ext in {".xlsx", ".xlsm"}:
            return self._extract_xlsx(path)
        if ext == ".csv":
            return path.read_text(encoding="utf-8", errors="replace")
        return (
            f"Uploaded document: {path.name}\n"
            "Asset URL: pending\n"
            f"MIME type: {mimetype or 'unknown'}\n"
            "This file type is stored as an attachment. Install a parser and re-ingest if full text extraction is required."
        )

    def _extract_pdf(self, path):
        try:
            from pypdf import PdfReader
        except Exception:
            try:
                from PyPDF2 import PdfReader
            except Exception:
                return self._unsupported_text(path, "PDF text extraction requires pypdf or PyPDF2.")
        try:
            reader = PdfReader(str(path))
            pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n\n".join(page.strip() for page in pages if page.strip())
            return text or self._unsupported_text(path, "No extractable text found in PDF.")
        except Exception as exc:
            return self._unsupported_text(path, f"PDF extraction failed: {exc}")

    def _extract_docx(self, path):
        try:
            from docx import Document
        except Exception:
            return self._unsupported_text(path, "DOCX extraction requires python-docx.")
        try:
            document = Document(str(path))
            paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    paragraphs.append(" | ".join(cell.text.strip() for cell in row.cells))
            return "\n".join(paragraphs) or self._unsupported_text(path, "No extractable text found in DOCX.")
        except Exception as exc:
            return self._unsupported_text(path, f"DOCX extraction failed: {exc}")

    def _extract_xlsx(self, path):
        try:
            from openpyxl import load_workbook
        except Exception:
            return self._unsupported_text(path, "XLSX extraction requires openpyxl.")
        try:
            workbook = load_workbook(str(path), read_only=True, data_only=True)
            lines = []
            for sheet in workbook.worksheets:
                lines.append(f"Sheet: {sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    values = [str(value) for value in row if value is not None]
                    if values:
                        lines.append(" | ".join(values))
            return "\n".join(lines) or self._unsupported_text(path, "No extractable values found in spreadsheet.")
        except Exception as exc:
            return self._unsupported_text(path, f"Spreadsheet extraction failed: {exc}")

    def _unsupported_text(self, path, reason):
        return (
            f"Uploaded document: {path.name}\n"
            "Asset URL: pending\n"
            f"Extraction note: {reason}\n"
            "The original file is linked to this memory through the tokenized asset URL."
        )

    def _chunk_text(self, text, chunk_size):
        text = text.strip()
        if not text:
            return ["Empty uploaded document."]
        chunks = []
        start = 0
        while start < len(text):
            end = min(start + chunk_size, len(text))
            if end < len(text):
                boundary = text.rfind("\n", start, end)
                if boundary > start + chunk_size // 2:
                    end = boundary
            chunks.append(text[start:end].strip())
            start = end
        return [chunk for chunk in chunks if chunk]
